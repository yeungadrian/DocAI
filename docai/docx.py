from typing import IO

import docx
from tabulate import tabulate

from docai.document import Document, MetaData
from docai.settings import Settings


def parse_docx(
    file: IO[bytes], filename: str, settings: Settings = Settings()
) -> list[Document]:
    documents = []
    for content in docx.Document(file).iter_inner_content():
        if isinstance(content, docx.table.Table):
            tabular_data = [[cell.text for cell in row.cells] for row in content.rows]
            documents.append(
                Document(
                    content=tabulate(
                        tabular_data,
                        tablefmt=settings.tables.tablefmt,
                        showindex=settings.tables.showindex,
                    ),
                    metadata=MetaData(filename=filename),
                )
            )
        elif isinstance(content, docx.text.paragraph.Paragraph):
            # TODO: Handle different styles, lists, headings etc
            documents.append(
                Document(
                    content=content.text,
                    metadata=MetaData(filename=filename),
                )
            )
    return documents
